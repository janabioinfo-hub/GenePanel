import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import os
import re

# Page config MUST be first
st.set_page_config(
    page_title="Gene Coverage Analyzer",
    page_icon="🧬",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Minimal CSS - only essential styles
st.markdown("""
    <style>
    .stButton>button {
        width: 100%;
    }
    </style>
""", unsafe_allow_html=True)

# Initialize session state ONCE
if 'coverage_data' not in st.session_state:
    st.session_state.coverage_data = None
if 'panel_genes' not in st.session_state:
    st.session_state.panel_genes = []
if 'filtered_data' not in st.session_state:
    st.session_state.filtered_data = None
if 'processed_csv' not in st.session_state:
    st.session_state.processed_csv = None
if 'file_basename' not in st.session_state:
    st.session_state.file_basename = "output"
if 'last_processed_file' not in st.session_state:
    st.session_state.last_processed_file = None
if 'mito_data' not in st.session_state:
    st.session_state.mito_data = None

@st.cache_data
def extract_gene_name(name):
    """Extract the first gene name from comma or semicolon-separated names."""
    if isinstance(name, str):
        if ',' in name:
            return name.split(',')[0]
        elif ';' in name:
            return name.split(';')[0]
        return name
    return None

def parse_gene_list(text):
    """Parse gene list with support for newlines, commas, and spaces. Remove duplicates."""
    if not text.strip():
        return []
    
    # Replace newlines and commas with spaces
    text = text.replace('\n', ' ').replace(',', ' ').replace('\t', ' ')
    
    # Split by whitespace and filter empty strings
    genes = [gene.strip() for gene in text.split() if gene.strip()]
    
    # Remove duplicates while preserving order
    seen = set()
    unique_genes = []
    for gene in genes:
        if gene not in seen:
            seen.add(gene)
            unique_genes.append(gene)
    
    return unique_genes

@st.cache_data(show_spinner=False)
def preprocess_excel_cached(file_bytes, file_name, skip_rows=1):
    """Cached preprocessing to avoid re-computation"""
    try:
        data = pd.read_excel(io.BytesIO(file_bytes), skiprows=skip_rows)
        
        if 'Gene Name' not in data.columns:
            raise ValueError("Column 'Gene Name' not found")
        
        data['Gene_Name'] = data['Gene Name'].apply(extract_gene_name)
        
        result = []
        identifiers = data[['Gene Names', 'Aliases', 'Gene_Name']].fillna('')
        unique_ids = identifiers.drop_duplicates()
        
        for _, row in unique_ids.iterrows():
            gene = row['Gene Names']
            alias = row['Aliases']
            name = row['Gene_Name']
            
            if gene:
                gene_data = data[data['Gene Names'] == gene]
            elif alias:
                gene_data = data[data['Aliases'] == alias]
            elif name:
                gene_data = data[data['Name'] == name]
            else:
                continue
            
            total_counted_bases = gene_data['Counted Bases'].sum()
            if total_counted_bases == 0:
                continue
            
            mean_depth = (gene_data['Mean Depth'] * gene_data['Counted Bases']).sum() / total_counted_bases
            mean_1x = (gene_data['% 1x'] * gene_data['Counted Bases']).sum() / total_counted_bases
            
            summary = {
                'Region': 'total',
                'Ref Name': gene if gene else None,
                'Aliases': alias if alias else None,
                'Gene_Name': name if name else None,
                'Name': gene_data['Name'].iloc[0] if 'Name' in gene_data.columns else None,
                'Gene IDs': gene_data['Gene IDs'].iloc[0] if 'Gene IDs' in gene_data.columns else None,
                'Counted Bases': total_counted_bases,
                'Mean Depth': mean_depth,
                'Min Depth': gene_data['Min Depth'].min(),
                'Max Depth': gene_data['Max Depth'].max(),
                '% 1x': mean_1x,
            }
            result.append(summary)
        
        coverage_df = pd.DataFrame(result)
        base_name = os.path.splitext(file_name)[0]
        
        return coverage_df, base_name
    except Exception as e:
        raise Exception(f"Error: {str(e)}")

def create_word_document(df_grouped, output_filename):
    """Create Word document with gene coverage table (legacy function - uses Perc_1x)"""
    # Rename for compatibility
    df_work = df_grouped.copy()
    if 'Perc_1x' in df_work.columns:
        df_work = df_work.rename(columns={'Perc_1x': '% 1x'})
    
    return create_word_document_with_mito(df_work, output_filename)

def create_word_document_with_mito(df_data, output_filename):
    """Create Word document with gene coverage table - works with mito or regular data"""
    
    # IMPORTANT: Reset the index to avoid iloc issues
    df_data = df_data.reset_index(drop=True)
    
    doc = Document()
    doc.add_heading('Appendix 1: Gene Coverage', 1)
    doc.add_heading('Indication Based Analysis:', 2)
    
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    
    chunk_size = 4
    chunks = [df_data[i:i+chunk_size] for i in range(0, len(df_data), chunk_size)]
    
    table = doc.add_table(rows=1, cols=chunk_size*2)
    table.alignment = 1
    hdr_cells = table.rows[0].cells
    
    for i in range(chunk_size):
        gene_hdr_cell = hdr_cells[i*2]
        gene_hdr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        gene_hdr_para = gene_hdr_cell.paragraphs[0]
        gene_hdr_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        gene_hdr_para.paragraph_format.space_after = Pt(0)
        gene_hdr_run = gene_hdr_para.add_run("Gene Name")
        gene_hdr_run.font.name = 'Calibri'
        gene_hdr_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        gene_hdr_run.font.size = Pt(8)
        
        perc_hdr_cell = hdr_cells[i*2+1]
        perc_hdr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        perc_hdr_para = perc_hdr_cell.paragraphs[0]
        perc_hdr_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        perc_hdr_para.paragraph_format.space_after = Pt(0)
        perc_hdr_run = perc_hdr_para.add_run("Percentage of coding region covered")
        perc_hdr_run.font.name = 'Calibri'
        perc_hdr_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        perc_hdr_run.font.size = Pt(8)
    
    for chunk in chunks:
        # IMPORTANT: Reset index for each chunk too
        chunk = chunk.reset_index(drop=True)
        
        row_cells = table.add_row().cells
        for i in range(chunk_size):
            if i < len(chunk):
                row = chunk.iloc[i]
                gene = str(row['Gene_ID'])
                
                # Get the percentage value - now using the correct index
                percent_value = row['% 1x']
                if isinstance(percent_value, pd.Series):
                    percent = float(percent_value.iloc[0])
                else:
                    percent = float(percent_value) if pd.notna(percent_value) else 0.0
                
                gene_cell = row_cells[i*2]
                gene_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                gene_para = gene_cell.paragraphs[0]
                gene_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                gene_para.paragraph_format.space_after = Pt(0)
                gene_run = gene_para.add_run(gene)
                gene_run.font.name = 'Calibri'
                gene_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                gene_run.font.size = Pt(8)
                gene_run.italic = True
                if percent < 90:
                    gene_run.font.color.rgb = RGBColor(255, 0, 0)
                
                perc_cell = row_cells[i*2+1]
                perc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                perc_para = perc_cell.paragraphs[0]
                perc_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                perc_para.paragraph_format.space_after = Pt(0)
                perc_run = perc_para.add_run(str(percent))
                perc_run.font.name = 'Calibri'
                perc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                perc_run.font.size = Pt(8)
                if percent < 90:
                    perc_run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                for empty_cell in [row_cells[i*2], row_cells[i*2+1]]:
                    empty_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    para = empty_cell.paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    para.paragraph_format.space_after = Pt(0)
                    run = para.add_run("–")
                    run.font.name = 'Calibri'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(255, 255, 255)
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    return doc

# Main UI
st.title("🧬 Gene Coverage Analyzer")
st.caption("Process raw Excel data or upload pre-processed CSV to generate Word report")

# Step 1
input_type = st.radio(
    "**Step 1: Choose Input Type**",
    ["📊 Pre-processed CSV", "📁 Raw Excel File"],
    horizontal=True
)

# Step 2
col1, col2 = st.columns(2)

with col1:
    if "Raw Excel" in input_type:
        st.subheader("Step 2A: Upload Raw Excel")
        raw_excel = st.file_uploader("Choose Excel file", type=['xlsx', 'xls'], key='raw_excel')
        
        if raw_excel:
            file_id = f"{raw_excel.name}_{raw_excel.size}"
            
            if st.session_state.last_processed_file != file_id:
                progress = st.progress(0, "Reading file...")
                
                try:
                    progress.progress(30, "Processing...")
                    file_bytes = raw_excel.read()
                    coverage_df, basename = preprocess_excel_cached(file_bytes, raw_excel.name)
                    
                    progress.progress(80, "Generating CSV...")
                    st.session_state.coverage_data = coverage_df
                    st.session_state.file_basename = basename
                    st.session_state.processed_csv = coverage_df.to_csv(index=False)
                    st.session_state.last_processed_file = file_id
                    
                    progress.progress(100, "Complete!")
                    progress.empty()
                except Exception as e:
                    progress.empty()
                    st.error(f"❌ {str(e)}")
            
            if st.session_state.processed_csv:
                st.success(f"✅ {len(st.session_state.coverage_data)} records")
                st.download_button(
                    "⬇️ Download CSV",
                    data=st.session_state.processed_csv,
                    file_name=f"{st.session_state.file_basename}_coverage.csv",
                    mime="text/csv",
                    use_container_width=True
                )
    else:
        st.subheader("Step 2A: Upload CSV Files")
        coverage_file = st.file_uploader("Choose Coverage CSV", type=['csv'], key='coverage')
        
        if coverage_file:
            st.session_state.file_basename = os.path.splitext(coverage_file.name)[0]
            df = pd.read_csv(coverage_file)
            df.columns = df.columns.str.strip()
            
            coverage_col = next((col for col in ['% 1x', '%1x', '% 1x '] if col in df.columns), None)
            
            if coverage_col:
                st.session_state.coverage_data = df
                st.success(f"✅ {len(df)} records")
            else:
                st.error("❌ No coverage column found")
        
        # Mitochondrial file upload option
        st.markdown("---")
        st.markdown("**Optional: Mitochondrial Data**")
        mito_file = st.file_uploader(
            "Upload Mitochondrial Excel (Optional)",
            type=['xlsx', 'xls'],
            key='mito_file',
            help="Excel file with mitochondrial gene coverage (header in row 2)"
        )
        
        if mito_file:
            try:
                # Load mito file with header in row 2 (skip first row)
                mito_df = pd.read_excel(mito_file, header=1)
                
                # Strip all column names of leading/trailing spaces
                mito_df.columns = mito_df.columns.str.strip()
                
                # Find %1x column (more flexible matching)
                mito_cols = [c for c in mito_df.columns if '1x' in str(c).lower()]
                
                if not mito_cols:
                    st.error("❌ No coverage column found in mitochondrial file")
                    st.write("Available columns:", mito_df.columns.tolist())
                    st.session_state.mito_data = None
                else:
                    mito_percent_col = mito_cols[0]
                    st.write(f"DEBUG - Using column: '{mito_percent_col}'")
                    
                    # Rename to standard name
                    mito_df = mito_df.rename(columns={mito_percent_col: '% 1x'})
                    
                    # Extract gene names from Name column (before '/')
                    if 'Name' not in mito_df.columns:
                        st.error("❌ No 'Name' column found in mitochondrial file")
                        st.write("Available columns:", mito_df.columns.tolist())
                        st.session_state.mito_data = None
                    else:
                        # Extract gene ID
                        mito_df['Gene_ID'] = mito_df['Name'].astype(str).str.split('/').str[0].str.strip()
                        
                        # Keep only needed columns
                        mito_df = mito_df[['Gene_ID', '% 1x']].copy()
                        
                        # Convert to numeric and round
                        mito_df['% 1x'] = pd.to_numeric(mito_df['% 1x'], errors='coerce')
                        
                        # Remove rows with NaN coverage
                        before_count = len(mito_df)
                        mito_df = mito_df.dropna(subset=['% 1x'])
                        after_count = len(mito_df)
                        
                        if before_count > after_count:
                            st.warning(f"⚠️ Removed {before_count - after_count} rows with invalid coverage values")
                        
                        # Round the values
                        mito_df['% 1x'] = mito_df['% 1x'].round(2)
                        
                        # Remove empty gene names
                        mito_df = mito_df[mito_df['Gene_ID'].str.strip() != '']
                        
                        # DEBUG: Show sample data
                        st.write(mito_df.head())
                        
                        st.session_state.mito_data = mito_df
                        st.success(f"✅ Loaded {len(mito_df)} mitochondrial genes")
                        
            except Exception as e:
                st.error(f"❌ Error loading mitochondrial file: {str(e)}")
                import traceback
                st.code(traceback.format_exc())
                st.session_state.mito_data = None
        else:
            st.session_state.mito_data = None

with col2:
    st.subheader("Step 2B: Panel Genes")
    tab1, tab2 = st.tabs(["📁 Excel", "📝 Paste"])
    
    with tab1:
        panel_file = st.file_uploader("Choose Excel", type=['xlsx', 'xls'], key='panel')
        if panel_file:
            panel_df = pd.read_excel(panel_file)
            if 'GENE' in panel_df.columns:
                genes = panel_df['GENE'].dropna().unique().tolist()
                st.session_state.panel_genes = genes
                st.success(f"✅ {len(genes)} genes")
            else:
                st.error("❌ No GENE column")
    
    with tab2:
        gene_text = st.text_area(
            "Paste genes (newline, comma, or space separated)",
            height=150,
            placeholder="BRCA1 BRCA2 TP53\nEGFR, KRAS\nTP53"
        )
        
        if st.button("Load Genes", use_container_width=True):
            genes = parse_gene_list(gene_text)
            if genes:
                st.session_state.panel_genes = genes
                duplicates_removed = len(gene_text.split()) - len(genes)
                if duplicates_removed > 0:
                    st.success(f"✅ {len(genes)} unique genes ({duplicates_removed} duplicates removed)")
                else:
                    st.success(f"✅ {len(genes)} genes")
            else:
                st.warning("⚠️ No genes found")

# Process data
if st.session_state.coverage_data is not None and st.session_state.panel_genes:
    df = st.session_state.coverage_data
    panel_genes = st.session_state.panel_genes
    
    coverage_col = next((col for col in ['% 1x', '%1x', '% 1x '] if col in df.columns), None)
    if not coverage_col:
        st.error("Coverage column not found")
        st.stop()
    
    df_filtered = df[df['Gene_Name'].isin(panel_genes) | df['Ref Name'].isin(panel_genes)].copy()
    df_filtered['Gene_ID'] = df_filtered['Gene_Name'].fillna(df_filtered['Ref Name'])
    df_filtered = df_filtered[~df_filtered['Gene_ID'].fillna('').str.startswith("Intron:")]
    df_filtered['Perc_1x'] = pd.to_numeric(df_filtered[coverage_col], errors='coerce')
    
    df_grouped = df_filtered.groupby('Gene_ID', as_index=False)['Perc_1x'].mean()
    df_grouped['Perc_1x'] = df_grouped['Perc_1x'].round(2)
    df_grouped = df_grouped.sort_values('Gene_ID')
    
    st.session_state.filtered_data = df_grouped
    
    st.divider()
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Genes", len(df_grouped))
    with col2:
        st.metric("Low Coverage", len(df_grouped[df_grouped['Perc_1x'] < 90]))
    with col3:
        st.metric("Avg Coverage", f"{df_grouped['Perc_1x'].mean():.2f}%")
    
    st.divider()
    
    col1, col2 = st.columns(2)
    

    with col1:
        if st.button("📄 Generate Word", type="primary", use_container_width=True):
            with st.spinner("Creating..."):
                try:
                    # Check if mito data exists
                    has_mito = 'mito_data' in st.session_state and st.session_state.mito_data is not None
                    
                    # Prepare final dataset
                    if has_mito:
                        # CRITICAL FIX: Rename panel genes column BEFORE concatenating
                        # df_grouped has column 'Perc_1x'
                        # mito_data has column '% 1x'
                        # They must match before concat!
                        
                        df_panel = df_grouped.copy()
                        df_panel = df_panel.rename(columns={'Perc_1x': '% 1x'})
                        df_panel = df_panel.sort_values('Gene_ID')
                        
                        # DEBUG: Verify column names match
                        
                        # Now both dataframes have the same column name: '% 1x'
                        df_final = pd.concat([
                            df_panel,
                            st.session_state.mito_data
                        ], ignore_index=True)
                        
                        # DEBUG: Check after concat
                        st.write(df_final.head(10))
                        
                        # Remove duplicates (keep first occurrence - panel genes take precedence)
                        df_final = df_final.drop_duplicates(subset='Gene_ID', keep='first')
                        
                        st.info(f"📊 Combined: {len(df_panel)} panel genes + {len(st.session_state.mito_data)} mito genes = {len(df_final)} total")
                    else:
                        # No mito data - just use panel data
                        df_final = df_grouped.copy()
                        df_final = df_final.rename(columns={'Perc_1x': '% 1x'})
                    
                    # Verify the final dataframe
                    st.write(df_final.head())
                    
                    word_filename = f"{st.session_state.file_basename}_report.docx"
                    doc = create_word_document_with_mito(df_final, word_filename)
                    doc_bytes = io.BytesIO()
                    doc.save(doc_bytes)
                    doc_bytes.seek(0)
                    st.session_state.doc_bytes = doc_bytes
                    st.session_state.word_filename = word_filename
                    st.success("✅ Ready!")
                except Exception as e:
                    st.error(f"❌ {str(e)}")
                    import traceback
                    st.code(traceback.format_exc())
        
        if 'doc_bytes' in st.session_state:
            st.download_button(
                f"⬇️ {st.session_state.word_filename}",
                data=st.session_state.doc_bytes,
                file_name=st.session_state.word_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )
    
    with col2:
        csv_filename = f"{st.session_state.file_basename}_filtered.csv"
        st.download_button(
            f"📊 {csv_filename}",
            data=df_grouped.to_csv(index=False),
            file_name=csv_filename,
            mime="text/csv",
            use_container_width=True
        )
    
    with st.expander("👁️ View Genes"):
        # Create styled HTML preview like before
        preview_html = '<div style="background: white; padding: 1rem; border-radius: 8px; border: 1px solid #dee2e6; max-height: 300px; overflow-y: auto;">'
        for _, row in df_grouped.iterrows():
            if row['Perc_1x'] < 90:
                preview_html += f'<span style="display: inline-block; padding: 0.3rem 0.6rem; margin: 0.2rem; background: #f8d7da; color: #721c24; border-radius: 4px; font-family: monospace; font-size: 0.9rem; font-weight: bold;"><i>{row["Gene_ID"]}</i> ({row["Perc_1x"]}%)</span>'
            else:
                preview_html += f'<span style="display: inline-block; padding: 0.3rem 0.6rem; margin: 0.2rem; background: #e9ecef; border-radius: 4px; font-family: monospace; font-size: 0.9rem;"><i>{row["Gene_ID"]}</i> ({row["Perc_1x"]}%)</span>'
        preview_html += '</div>'
        st.markdown(preview_html, unsafe_allow_html=True)
        st.caption("ℹ️ Genes with coverage below 90% are highlighted in red")

st.divider()

# Batch HTML Generator
with st.expander("📦 Batch HTML Report Generator"):
    st.markdown("""
    **Generate interactive HTML reports from multiple CSV files**
    
    Upload multiple CSV files and get a ZIP with beautiful HTML reports for each file.
    Each report includes search, sorting, and download features.
    """)
    
    batch_files = st.file_uploader(
        "Upload Multiple CSV Files",
        type=['csv'],
        accept_multiple_files=True,
        key='batch_csvs',
        help="Select all CSV files to convert to HTML"
    )
    
    if batch_files:
        st.info(f"✅ Loaded {len(batch_files)} files")
        
        if st.button("🎨 Generate HTML Reports", type="primary", use_container_width=True):
            progress_bar = st.progress(0, f"Generating reports...")
            
            try:
                import zipfile
                from datetime import datetime
                
                zip_buffer = io.BytesIO()
                
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for idx, csv_file in enumerate(batch_files):
                        df = pd.read_csv(csv_file)
                        patient_name = os.path.splitext(csv_file.name)[0]
                        html_content = generate_html_report(df, patient_name)
                        
                        safe_filename = patient_name.replace(' ', '_').replace('/', '_')
                        zip_file.writestr(f"coverage_{safe_filename}.html", html_content)
                        
                        progress = int((idx + 1) / len(batch_files) * 100)
                        progress_bar.progress(progress / 100, f"Creating {idx + 1} of {len(batch_files)}...")
                
                zip_buffer.seek(0)
                progress_bar.empty()
                
                st.success(f"✅ Generated {len(batch_files)} reports!")
                
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                st.download_button(
                    label=f"⬇️ Download All Reports (ZIP)",
                    data=zip_buffer.getvalue(),
                    file_name=f"gene_coverage_reports_{timestamp}.zip",
                    mime="application/zip",
                    use_container_width=True,
                    type="primary"
                )
                
            except Exception as e:
                progress_bar.empty()
                st.error(f"❌ Error: {str(e)}")

def generate_html_report(df, patient_name):
    """Generate interactive HTML report"""
    
    # Get gene column name
    gene_col = 'Gene_ID' if 'Gene_ID' in df.columns else 'Gene_Name' if 'Gene_Name' in df.columns else df.columns[0]
    coverage_col = 'Perc_1x' if 'Perc_1x' in df.columns else 'Coverage' if 'Coverage' in df.columns else df.columns[1]
    
    rows_html = []
    for idx, row in df.iterrows():
        gene = row.get(gene_col, '')
        coverage = row.get(coverage_col, '')
        rows_html.append(f"<tr><td>{idx + 1}</td><td>{gene}</td><td>{coverage}</td></tr>")
    
    table_rows = '\n'.join(rows_html)
    avg_coverage = df[coverage_col].mean() if coverage_col in df.columns else 0
    
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Gene Coverage Report - {patient_name}</title>
  <style>
    body {{
      margin: 0;
      font-family: Arial, sans-serif;
      background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
      padding: 20px;
    }}
    .container {{
      max-width: 1200px;
      margin: 0 auto;
      background: white;
      border-radius: 15px;
      padding: 30px;
      box-shadow: 0 20px 60px rgba(0,0,0,0.3);
    }}
    .header {{
      text-align: center;
      margin-bottom: 30px;
      padding-bottom: 20px;
      border-bottom: 3px solid #667eea;
    }}
    .header h1 {{
      color: #667eea;
      margin: 0;
      font-size: 2em;
    }}
    .patient-name {{
      color: #666;
      font-size: 1.2em;
      margin-top: 10px;
    }}
    .controls {{
      margin: 20px 0;
      display: flex;
      gap: 10px;
      flex-wrap: wrap;
      justify-content: center;
    }}
    input[type="text"] {{
      padding: 10px 15px;
      font-size: 14px;
      border: 2px solid #667eea;
      border-radius: 8px;
      min-width: 300px;
    }}
    button {{
      padding: 10px 20px;
      background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
      color: white;
      border: none;
      border-radius: 8px;
      cursor: pointer;
      font-size: 14px;
      font-weight: 600;
      transition: transform 0.2s;
    }}
    button:hover {{
      transform: translateY(-2px);
      box-shadow: 0 5px 15px rgba(102, 126, 234, 0.4);
    }}
    .table-wrapper {{
      max-height: 600px;
      overflow-y: auto;
      border: 2px solid #667eea;
      border-radius: 10px;
      margin-top: 20px;
    }}
    table {{
      border-collapse: collapse;
      width: 100%;
    }}
    th, td {{
      border: 1px solid #dee2e6;
      padding: 12px;
      text-align: center;
      font-size: 14px;
    }}
    th {{
      background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
      color: white;
      position: sticky;
      top: 0;
      font-weight: 600;
      z-index: 10;
    }}
    tr:nth-child(even) {{
      background-color: #f8f9fa;
    }}
    tr:hover {{
      background-color: #e9ecef;
    }}
    .stats {{
      display: flex;
      gap: 20px;
      justify-content: center;
      margin: 20px 0;
    }}
    .stat-box {{
      background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
      color: white;
      padding: 15px 30px;
      border-radius: 10px;
      text-align: center;
    }}
    .stat-box h3 {{
      margin: 0;
      font-size: 2em;
    }}
    .stat-box p {{
      margin: 5px 0 0 0;
      opacity: 0.9;
    }}
    @media only screen and (max-width: 600px) {{
      .table-wrapper {{ max-height: 400px; }}
      table {{ font-size: 12px; }}
      .stats {{ flex-direction: column; }}
    }}
  </style>
</head>
<body>
  <div class="container">
    <div class="header">
      <h1>🧬 Gene Coverage Report</h1>
      <div class="patient-name">{patient_name}</div>
    </div>

    <div class="stats">
      <div class="stat-box">
        <h3 id="totalGenes">{len(df)}</h3>
        <p>Total Genes</p>
      </div>
      <div class="stat-box">
        <h3 id="avgCoverage">{avg_coverage:.2f}%</h3>
        <p>Average Coverage</p>
      </div>
    </div>

    <div class="controls">
      <input type="text" id="searchInput" placeholder="🔍 Search by Gene..." />
      <button onclick="sortTable(true)">↑ Ascending</button>
      <button onclick="sortTable(false)">↓ Descending</button>
      <button onclick="downloadCSV()">📥 Download CSV</button>
    </div>

    <div class="table-wrapper">
      <table id="dataTable">
        <thead>
          <tr>
            <th>S.No</th>
            <th>Gene</th>
            <th>Coverage (%)</th>
          </tr>
        </thead>
        <tbody>
          {table_rows}
        </tbody>
      </table>
    </div>
  </div>

  <script>
    const searchInput = document.getElementById('searchInput');
    const tbody = document.querySelector('tbody');
    let rows = Array.from(tbody.querySelectorAll('tr'));

    searchInput.addEventListener('keyup', function(e) {{
      const searchText = e.target.value.toLowerCase();
      rows.forEach(row => {{
        const found = Array.from(row.cells).some(cell => 
          cell.textContent.toLowerCase().includes(searchText)
        );
        row.style.display = found ? '' : 'none';
      }});
      updateStats();
    }});

    function sortTable(ascending) {{
      const sortedRows = rows.slice().sort((a, b) => {{
        const aVal = parseFloat(a.cells[2].textContent) || 0;
        const bVal = parseFloat(b.cells[2].textContent) || 0;
        return ascending ? aVal - bVal : bVal - aVal;
      }});
      sortedRows.forEach((row, i) => {{
        row.cells[0].textContent = i + 1;
        tbody.appendChild(row);
      }});
      rows = sortedRows;
    }}

    function updateStats() {{
      const visible = rows.filter(row => row.style.display !== 'none');
      document.getElementById('totalGenes').textContent = visible.length;
    }}

    function downloadCSV() {{
      const visible = rows.filter(row => row.style.display !== 'none');
      let csv = "S.No,Gene,Coverage\\n";
      visible.forEach(row => {{
        const cells = Array.from(row.cells).map(c => '"' + c.textContent.trim() + '"');
        csv += cells.join(',') + '\\n';
      }});
      const blob = new Blob([csv], {{ type: 'text/csv' }});
      const link = document.createElement('a');
      link.href = URL.createObjectURL(blob);
      link.download = 'coverage_{patient_name}.csv';
      link.click();
    }}
  </script>
</body>
</html>"""
    
    return html

st.divider()
st.caption("All processing happens on the server. Data is not stored permanently.")
